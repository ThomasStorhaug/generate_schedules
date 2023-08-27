import locale
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.run import Run

import settings

"""
time: "kode:fag:rom:lærer", ved fri: "kode:beskrivelse", eks: "3:Grunnlovsdagen", 
    koder:
        "1" - enkelttime
        "2" - dobbelttime
        "3" - fridag
timeplan: [
[1.time, 1.time, 1.time,...],
[2.time, 2.time, ...],
[3.time,...],
...
]
"""

def convert_to_rgb(color_hex)->RGBColor:
    """
    :param color_hex: string: #RRGGBB

    return RGBColor() instance
    """
    r = int(color_hex[1:3], 16)
    g = int(color_hex[3:5], 16)
    b = int(color_hex[5:7], 16)

    return RGBColor(r, g, b)

def shade_cell(cell: _Cell, color_str: str):
    """
    Word XML magic
    :param cell: the _Cell object that will get shaded
    :param color_str: a string representation of the rgb color, can be with or without #
    """
    cell_element = cell._tc

    shading_elm = OxmlElement("w:shd")

    shading_elm.set(qn("w:fill"), color_str)

    cell_element.get_or_add_tcPr().append(shading_elm)

def set_vertical_alignment(cell: _Cell, align:str="center"):
    """
    Sets the vertical alignment of a cell, default is center aligned
    :param cell: the _Cell object of which alignment will be set
    :param align: can be top, center og bottom
    """

    if align not in('top', 'center', 'bottom'):
        raise ValueError("Invalid alignment value.")
    
    cell_element = cell._tc

    cell_properties = cell_element.get_or_add_tcPr()

    vAlign_elm = OxmlElement('w:vAlign')
    vAlign_elm.set(qn('w:val'), align)

    cell_properties.append(vAlign_elm)

def set_horizontal_alignment(cell:_Cell, align:WD_ALIGN_PARAGRAPH):
    """
    Set horizontal alignment of text in a table cell.

    :param cell: a cell in a Word table
    :param align: a member of the WD_ALIGN_PARAGRAPH enumeration
    """
    for paragraph in cell.paragraphs:
        paragraph.alignment = align

def set_text_color(run, color):
    """
    Sets the color of the text in a run
    :param run: a text run of a paragraph
    :param color: a color string
    """

    run.font.color.rgb = convert_to_rgb(color)

def insert_text_in_cell(cell:_Cell, text:str, alignment:WD_ALIGN_PARAGRAPH=None, font:str="Calibri", size=11)->Run:
    """
    Inserts text in a cell, returns the run containing the text
    :param cell: the cell where the text will be inserted
    :param text: the text to insert
    :param alignment: optioinal, must be a WD_ALIGN_PARAGRAPH type
    :param font: optional, defaults to Calibri
    :param size: optional, defaults to 11
    """
    if alignment != None:
        set_horizontal_alignment(cell, alignment)

    run = cell.paragraphs[0].add_run(text)
    run.font.name = font
    run.font.size = Pt(size)

    return run

def create_base_table(doc:Document) -> Table:
    """
    Creates the base table of the schedule. This includes inserting text in the first column (times) and
    setting the height and width of cells. Returns the table object
    :param doc: the Document object
    """
    hori_align = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=13, cols=6)
    table.autofit = False

    class_rows = [table.rows[2], table.rows[3], table.rows[5], table.rows[6], table.rows[8], table.rows[9], table.rows[11], table.rows[12]]

    for cell in table.columns[0].cells:
        cell.width = settings.TIMES_COLUMN_WIDTH

    for column in list(table.columns)[1:]:
        for cell in column.cells:
            cell.width = settings.CLASS_COLUMN_WIDTH
    
    for row in class_rows:
        row.height = settings.CLASS_CELL_HEIGHT

    
    top_row = table.cell(0, 0).merge(table.cell(0, 5))
    table.style = "Table Grid"
    top_row.text = settings.TOP_ROW_TEXT
    
    pause_cells = {
        1: table.cell(4, 0),
        2: table.cell(7, 0),
        3: table.cell(10, 0)
    }

    # insert text in first column, pauses and periods.
    for key in pause_cells:
        insert_text_in_cell(pause_cells[key], f"Pause\n{settings.PAUSES[key]}", hori_align, size=9)
        shade_cell(pause_cells[key], settings.COLORS["accent"]["background"])
        set_vertical_alignment(pause_cells[key], "center")

    for cell in table.rows[1].cells:
        shade_cell(cell, settings.COLORS["accent"]["background"])

    for i, y in enumerate([2, 3, 5, 6, 8, 9, 11, 12]):
        insert_text_in_cell(table.cell(y, 0), f"{i+1}\n{settings.CLASS_START[i+1]}", hori_align, size=9)
        set_vertical_alignment(table.cell(y, 0))

    return table

def set_to_landscape(doc:Document):
    """
    Sets the documents first sections orientation to landscape
    :param doc: the Document object
    """
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

def set_margins(doc:Document, left:Cm, right:Cm):
    """
    Sets the margins of the documents first section
    :param doc: the Document object
    :param left: the left margin
    :param right: the right margin
    """
    section = doc.sections[0]
    section.right_margin = right
    section.left_margin = left

def get_date_string(date:datetime, language:str="no_NO"):
    """
    Converts datetime object to a date string formatted for the schedule
    :param date: datetime object
    :param language: optional for translating the schedule
    """
    locale.setlocale(locale.LC_TIME, language)
    return date.strftime("%A %d.%b").capitalize()

def insert_dates(table:Table, date_range:list):
    """
    Inserts dates in the second row of the table, starting at the second column

    :param table: table object
    :param date_range: list of datetime objects representing each work day of that week
    """
    for i, date in enumerate(date_range):
        insert_text_in_cell(table.cell(1, i+1), get_date_string(date), WD_ALIGN_PARAGRAPH.CENTER)

def insert_periods(table:Table, day:int, periods):
    """
    Inserts the schedule for an entire day.
    :param table: the Table object 
    :param periods: a list of periods, or a string for holidays, each represented as a string, see README for formatting
    :param day: integer representation of weekday, 1 = monday etc.
    :param holiday: 
    """
    y_coords = [2, 3, 5, 6, 8, 9, 11, 12]

    if type(periods) == str:
        # Holiday
        holiday = periods.split(":")[1]
        cell = table.cell(2, day).merge(table.cell(12, day))
        text_run = insert_text_in_cell(cell, holiday, WD_ALIGN_PARAGRAPH.CENTER, size = 9)
        shade_cell(cell, settings.COLORS["fri"]["background"])
        set_text_color(text_run, settings.COLORS["fri"]["text"])
        set_vertical_alignment(cell)
        return
    
    for period, y in zip(periods, y_coords):
        period_ls = period.split(":")

        if period_ls[0] == "2":
            try:
                cell = table.cell(y, day).merge(table.cell(y+1, day))
            except IndexError:
                pass
        else:
            cell = table.cell(y, day)
        
        if period_ls[1] == "x" or period_ls[1] == "-":
            continue
        if len(period_ls) > 2:
            text_str = ""
            for item in period_ls[2:]:
                text_str += "\n" + item
            code = f'{settings.SUBJECT_CODES[period_ls[1]]}'
            text = code + text_str
        else:
            text = f'{settings.SUBJECT_CODES[period_ls[1]]}'
            code = text
        text_run = insert_text_in_cell(cell, text, WD_ALIGN_PARAGRAPH.CENTER, size=9)
        shade_cell(cell, settings.COLORS[code]["background"])
        set_text_color(text_run, settings.COLORS[code]["text"])
        set_vertical_alignment(cell)

        
def parse_timetable(timetable:list):
    """
    Converts a list in a schedule format, where each secondary list is a table row, 
    to a list where each secondary list is a day.
    Holidays must follow this format : "3:explanation",
    Example: "3:Grunnlovsdagen"

    :param timetable: must be in a weekly schedule format
    """
    
    days = []
    has_a_day_off = False

    for i, subj in enumerate(timetable[0]):
        if subj.split(":")[0] == "3":
            days.append(subj)
            has_a_day_off = True
            continue
        else:
            day = []
            for j in range(8):
                day.append(timetable[j][i])

            days.append(day)
    if has_a_day_off:
        print(days)
    return days

def create_schedule(timetable:list, dates:datetime, path:str):
    """
    Creates and saves a word document at the given path with a full time schedule

    See README for extensive explanation of timetable formating.

    :param timetable: a two dimensional array of lists; primary list is the entire schedule, secondary lists represents rows in the table. Each school hour is represented as strings in the secondary list
    :param dates: a list of datetime objects representing 5 work days.
    :param path: the full path the file will be saved at.
    :param holidays: dict of {datetime.date:"name"} for each holiday in the given week
    """
    document = Document()
    set_to_landscape(document)
    set_margins(document, settings.LEFT_MARGIN, settings.RIGHT_MARGIN)
    table = create_base_table(document)

    days = timetable
    for i, day in enumerate(days):
        insert_periods(table, i+1, day)
    # Insert dates
    insert_dates(table, dates)
    document.save(path)

